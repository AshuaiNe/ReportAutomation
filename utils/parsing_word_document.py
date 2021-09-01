#!/usr/bin/env python
# encoding: utf-8
import re
import time
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shared import RGBColor
from win32com import client as wc
from utils.log import LoggerFactory
from shutil import copyfile
from utils.glo import Globals
from pathlib import Path
from threading import Thread, Lock


class ParsingWord:
    log = LoggerFactory(level="info")
    log.create_logger()

    def __init__(self) -> None:
        self.glo = Globals()
        self.compile = re.compile(r"[\n\s*\r]+")
        self._path = self.glo.get_value('_path')
        self.parsing = self.glo.get_value('parsing')
        self.product = self.glo.get_value('product')
        self.original_data = f"{self._path}/compare/original_data/{self.parsing}/{time.strftime('%Y-%m-%d', time.localtime())}/"
        self.challenger = f"{self._path}/compare/challenger/{self.parsing}/{time.strftime('%Y-%m-%d', time.localtime())}/"
        self._path_files = [self.original_data, self.challenger]
        self._result_file = f"{self._path}/compare/result/{self.parsing}/{time.strftime('%Y-%m-%d', time.localtime())}/"

    def iter_block_items(self, parent):
        '''
        @description: 迭代获取word文档中的段落和表格
        @param {*}
        @return {*}
        '''        
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        else:
            raise ValueError("something's not right")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def get_word(self, filename):
        '''
        @description: 重新组装获取的word结果
        @param {*}
        @return {*}
        '''        
        document = Document(filename)
        par_lists = []
        paragraph_data = {}
        table_data = {}
        set_title = ""
        title_table_SUM = 0
        x = True
        header_data = {}
        for block in self.iter_block_items(document):
            if isinstance(block, Paragraph):
                line = self.compile.sub("", block.text)
                if not line:
                    continue
                elif 'toc' in block.style.name:
                    continue
                elif "XBRL" not in block.style.name:
                    par_lists.append(line)
                elif "XBRL" in block.style.name:
                    if par_lists and x:
                        paragraph_data['cover'] = par_lists
                        x = False
                    elif not par_lists:
                        paragraph_data[set_title] = par_lists
                    else:
                        paragraph_data[set_title] = par_lists
                    set_title = line
                    par_lists = []
                    title_table_SUM = 0
            else:
                cell_dict = {}
                try:
                    if set_title not in table_data.keys():
                        table_data[set_title] = {}
                    for row_int, row in enumerate(block.rows):
                        B = []
                        for cel in row.cells:
                            line = re.sub(r"\s+", "", cel.text)
                            B.append(line)
                        if B[0] in cell_dict.keys():
                            B[0] = f"{B[0]}_{row_int}"
                        cell_dict.update({f"{B[0]}": B})
                        row_int += 1
                    title_table_SUM = len(table_data[set_title]) + title_table_SUM
                    table_title = set_title + f"_表格{title_table_SUM}"
                    table_data[set_title][table_title] = cell_dict
                    title_table_SUM += 1
                except Exception as e:
                    if set_title not in table_data.keys():
                        table_data[set_title] = {}
                    title_table_SUM = len(table_data[set_title]) + title_table_SUM
                    table_title = set_title + f"_表格{title_table_SUM}"
                    table_data[set_title][table_title] = {'异常表格，占位': '0'}
                    title_table_SUM += 1
                    self.log.logger.info(f'文件路径{filename}，标题：{set_title}表格取值异常，无法读取：{e}')
                    continue 
        if par_lists:
            paragraph_data[set_title] = par_lists
        if len(paragraph_data.keys()) < 3:
            self.log.logger.info(f"标题少于3个，格式异常：{filename}")
        if document.sections:
            for sec in document.sections:
                for par in sec.header.paragraphs:
                    header_data = par.text
        return paragraph_data, table_data, header_data


    def convert_doc_to_docx(self, app) -> str:
        '''
        @description: doc文件另存为docx
        @param {None}
        @return {self._path_files}
        '''    
        try:
            if app == 'OFFICE':
                application_app = "Word"
            elif app == 'WPS':
                application_app = "Kwps"
            try:
                w = wc.Dispatch(f'{application_app}.Application') # 打开word应用程序
            except Exception:
                self.log.logger.info(f"应用程序打开异常，请使用其它程序{Exception}")
            for path_file in self._path_files:
                for file in Path(path_file).iterdir():
                    if Path(file).suffix == '.doc' and not Path(file).stem.startswith('~$'):
                        doc = w.Documents.Open(file.__str__())
                        if app == 'OFFICE':
                            doc.SaveAs(f"{path_file}{Path(file).stem}" + '.docx', 12)
                        elif app == 'WPS':
                            doc.SaveAs2(f"{path_file}{Path(file).stem}" + '.docx', 12)
                        doc.Close()
                    else:
                        continue
            w.Quit()
        except Exception:
            self.log.logger.info(f"转换异常：{Exception}")
            pass


    def get_file_path_tuple(self, endswith):
        '''
        @description: 获取文件路径元组
        @param {*}
        @return {tuple}
        '''        
        try:
            original_file_list = []
            challenger_file_list = []
            i = 0
            for path_file in self._path_files:
                if Path(path_file).iterdir():
                    for file in Path(path_file).iterdir():
                        if Path(file).suffix == endswith and not Path(file).stem.startswith('~$'):
                            file_path = file.__str__()
                            if i == 0:
                                original_file_list.append(file_path)
                            elif i == 1 and endswith == '.xls':
                                challenger_file_list.append(file_path)
                            else:
                                file = "copyfile" + Path(file).stem + Path(file).suffix
                                file_new_path = self._result_file + file
                                copyfile(file_path, file_new_path).encode('utf-8')
                                challenger_file_list.append(file_new_path)
                        elif Path(file).suffix == '.doc':
                            continue
                        else:
                            self.log.logger.info(f"文件夹空，请核对\n文件内容{self._path_files}")
                else:
                    self.log.logger.info(f"文件夹空，请核对\n文件内容{self._path_files}")
                i = i + 1
            if len(original_file_list) != len(challenger_file_list):
                    self.log.logger.info(f"新旧文件数量不一致，请核对！\n文件内容{self._path_files}")
            else:
                return list(zip(sorted(original_file_list, reverse=False), sorted(challenger_file_list, reverse=False)))
        except Exception as e:
            self.log.logger.info(f"{e}")

    def set_docx(self, file_name, table_list, table_dict, compare_result):
        '''
        @description: 把比对结果写入word文档中
        @param {file_name, compare_result}
        @return {None}
        '''        
        paragraphs_old_value = {}
        title_old_value = []
        difference_title_list = []
        difference_paragraphs_list = {}
        difference_tables_list = {}
        tables_old_value = {}
        set_title = 'cover'
        doc = Document(file_name)
        for x, y, z in self.dict_generator(compare_result):
            if isinstance(z, dict):
                key = y
                new_value = re.sub(r"\s+", "", z['new_value'])
                old_value = re.sub(r"\s+", "", z['old_value'])
            else:
                key = z
            try:
                find_str = re.findall(r"\[.*?\]", key)
                tmp_str = find_str[0].replace("[", "").replace("]", "").replace("'", "")
            except:
                continue
            if 'dictionary_item_added' == x:
                    if tmp_str == 'docx_paragraphs':
                        difference_title_list.append(find_str[1].replace("[", "").replace("]", "").replace("'", ""))
            elif 'dictionary_item_removed' == x:
                    if tmp_str == 'docx_paragraphs':
                        title_old_value.append(find_str[1].replace("[", "").replace("]", "").replace("'", ""))
            elif 'values_changed' == x:
                    if 'docx_paragraphs' == tmp_str:
                        tmp_key = find_str[1].replace("[", "").replace("]", "").replace("'", "")
                        difference_paragraphs_list.update({f'{tmp_key}': f'{new_value}'})
                        paragraphs_old_value.update({f'{tmp_key}': f'{old_value}'})
                    elif 'docx_tables' == tmp_str:
                        difference_tables_list.update({f"{key}": new_value})
                        tables_old_value.update({f"{key}": old_value})
        for key, value in difference_paragraphs_list.items():
            for par in doc.paragraphs:
                if key != 'cover':
                    if "XBRL" in par.style.name:
                        set_title = re.sub(r"\s+", "", par.text)
                        continue
                line = re.sub(r"\s+", "", par.text)
                if line == value:
                    try:
                        if set_title == key:
                            for run in par.runs:
                                run.font.color.rgb = RGBColor(255, 255, 0)
                            par.add_comment(paragraphs_old_value[key], author='差异段落', initials='od')
                            break
                    except Exception as e:
                        continue
        for difference in difference_title_list:
            for par in doc.paragraphs:
                line = re.sub(r"\s+", "", par.text)
                if line == difference:
                    try:
                        for run in par.runs:
                            run.font.color.rgb = RGBColor(0, 255, 0)
                        par.add_comment('历史文件不存在该章节，请人工核对', author='新增章节', initials='od')
                        break
                    except Exception as e:
                        continue
        for key, value in difference_tables_list.items():
            find_str = re.findall(r"\[.*?\]", key)
            _table_str = str(find_str[2].replace("[", "").replace("]", "").replace("'", ""))
            _table = table_list.index(_table_str)
            _row = find_str[3].replace("[", "").replace("]", "").replace("'", "")
            _cel = int(find_str[4].replace("[", "").replace("]", "").replace("'", ""))
            _row_cel = [x for x in table_dict[_table_str][_row].keys()]
            _cel = int(_row_cel[0].split("_")[1]) + int(_cel)
            cell = doc.tables[int(_table)]._cells[_cel]
            line = re.sub(r"\s+", "", cell.text)
            if line == value:
                runs = cell.paragraphs[0].runs
                for run in runs:
                    run.font.color.rgb = RGBColor(255, 255, 0)
                cell.paragraphs[0].add_comment(tables_old_value[key], author='差异单元格', initials='od')
        for old_value in title_old_value:
            for x in range(10):
                if not doc.paragraphs[x].runs:
                    pass
                else:
                    run = doc.paragraphs[x].runs[-1]
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    run.add_comment(old_value.split("_")[0], author='删除的章节或章节下的表格', initials='od')
                    break
        doc.save(file_name)

    def dict_generator(self, indict):
        for key, value in indict.items():
            if isinstance(value, dict):
                for x, y in value.items():
                    yield key, x, y
            else:
                for _ in value:
                    yield key, None, _
